"""
doc_tools.py — Đọc/ghi DOCX cho việc cấu hình placeholder.

Chức năng:
  - scan_placeholders : quét mọi ${...} có trong file DOCX
  - generate_template : sinh DOCX template mới từ schema eform
  - inject_placeholders: chèn ${key} vào DOCX có sẵn theo bảng ánh xạ
"""

import re
from typing import Dict, List, Set

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from eform_parser import Field, GRID_TYPES, MULTI_CHECKBOX_TYPES

PLACEHOLDER_RE = re.compile(r"\$\{[^}]+\}")

NAVY = RGBColor(0x0D, 0x47, 0xA1)
GREY = RGBColor(0x54, 0x6E, 0x7A)
SHADE_INPUT = "F5F7FA"
SHADE_LABEL = "ECEFF1"


# ---------------------------------------------------------------- scan

def _iter_paragraphs(doc):
    """Duyệt mọi paragraph kể cả trong table lồng nhau."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def scan_placeholders(docx_path: str) -> Set[str]:
    """Trả về tập placeholder ${...} tìm thấy trong DOCX."""
    doc = Document(docx_path)
    found: Set[str] = set()
    for p in _iter_paragraphs(doc):
        found.update(PLACEHOLDER_RE.findall(p.text))
    # header/footer
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                found.update(PLACEHOLDER_RE.findall(p.text))
    return found


# ---------------------------------------------------------------- helpers

def _shade(cell, hex_fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_text(cell, text: str, bold=False, italic=False, size=11, color=None,
              align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_heading(doc, text: str, size=13, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = NAVY
    return p


def _add_label_value_row(doc, label: str, placeholder: str):
    """Bảng 2 cột: nhãn (nền xám) | placeholder (nền nhạt)."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    label_cell, value_cell = table.rows[0].cells
    _set_text(label_cell, label, bold=True, size=11)
    _shade(label_cell, SHADE_LABEL)
    _set_text(value_cell, placeholder, size=11)
    _shade(value_cell, SHADE_INPUT)
    doc.add_paragraph()
    return table


def _add_checkbox_table(doc, rows: List[tuple]):
    """Bảng 2 cột: nội dung | ô tick placeholder."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, placeholder) in enumerate(rows):
        text_cell, box_cell = table.rows[i].cells
        _set_text(text_cell, label, size=10.5)
        _set_text(box_cell, placeholder, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(box_cell, SHADE_INPUT)
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------- generate

def generate_template(fields: List[Field], out_path: str,
                      title: str = "BIỂU MẪU", subtitle: str = ""):
    """Sinh DOCX template mới, mọi field đều có placeholder đúng cú pháp."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Tiêu đề
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = GREY

    doc.add_paragraph()

    grid_keys = {f.key for f in fields if f.type in GRID_TYPES}
    subfields_by_grid: Dict[str, List[Field]] = {}
    for f in fields:
        if f.parent in grid_keys:
            subfields_by_grid.setdefault(f.parent, []).append(f)

    for f in fields:
        if f.parent in grid_keys:
            continue  # đã gộp vào bảng của grid

        if f.type in GRID_TYPES:
            _add_heading(doc, f.label or f.key)
            subs = subfields_by_grid.get(f.key, [])
            if subs:
                note = doc.add_paragraph()
                r = note.add_run("Cột: " + ", ".join(s.key for s in subs))
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = GREY
            ph = doc.add_paragraph()
            ph.add_run(f.placeholder).font.size = Pt(11)
            doc.add_paragraph()
            continue

        if f.type in MULTI_CHECKBOX_TYPES:
            _add_heading(doc, f.label or f.key)
            rows = [(lbl or val, ph) for val, lbl, ph in f.option_placeholders()]
            if rows:
                _add_checkbox_table(doc, rows)
            continue

        if f.type == "checkbox":
            _add_checkbox_table(doc, [(f.label or f.key, f.placeholder)])
            continue

        _add_label_value_row(doc, f.label or f.key, f.placeholder)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------- inject

def inject_placeholders(docx_path: str, out_path: str,
                        mapping: Dict[str, str]) -> int:
    """
    Chèn placeholder vào DOCX có sẵn.
    mapping: { "đoạn text cần thay": "${Key}" }
    Trả về số lần thay thế thành công.
    """
    doc = Document(docx_path)
    count = 0

    for p in _iter_paragraphs(doc):
        original = p.text
        replaced = original
        for needle, placeholder in mapping.items():
            if needle in replaced:
                replaced = replaced.replace(needle, placeholder)
        if replaced != original:
            # Ghi đè: dồn text vào run đầu, xoá các run còn lại
            for run in p.runs[1:]:
                run.text = ""
            if p.runs:
                p.runs[0].text = replaced
            else:
                p.add_run(replaced)
            count += 1

    doc.save(out_path)
    return count
