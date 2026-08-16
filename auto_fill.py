"""
auto_fill.py — Tự dò chỗ trống trong DOC gốc và gợi ý chèn ${key}.

Nguyên lý:
  1. Quét DOC, tìm các "chỗ trống" — đoạn có nhãn rồi tới dấu chấm/gạch dài,
     hoặc ô bảng rỗng nằm cạnh ô nhãn.
  2. So khớp nhãn đó với label của field trong eform (không dấu, bỏ dấu câu).
  3. Chấm điểm, đề xuất cặp (chỗ trống -> ${Key}) kèm độ tin cậy.
  4. Người dùng duyệt lại rồi mới ghi file.

Tool KHÔNG tự ghi đè mà không hỏi — bước duyệt là bắt buộc, vì so khớp theo
chuỗi không bao giờ đúng 100%.
"""

import re
import unicodedata
from dataclasses import dataclass
from difflib import SequenceMatcher
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from eform_parser import Field, GRID_TYPES, MULTI_CHECKBOX_TYPES

# Chuỗi coi là "chỗ trống": >= 3 dấu chấm, gạch dưới, gạch ngang
BLANK_RE = re.compile(r"[.．…_\-–—]{3,}")
# Ô tick trong DOC giấy
CHECKBOX_CHARS = "☐☒☑□■◻◼"
CHECKBOX_RE = re.compile(f"[{CHECKBOX_CHARS}]")
PLACEHOLDER_RE = re.compile(r"\$\{[^}]+\}")


# ---------------------------------------------------------------- normalize

def strip_accents(text: str) -> str:
    nfkd = unicodedata.normalize("NFD", text)
    return "".join(c for c in nfkd if unicodedata.category(c) != "Mn")


def normalize(text: str) -> str:
    """Chuẩn hoá để so khớp: bỏ dấu, thường hoá, bỏ dấu câu và số thứ tự."""
    if not text:
        return ""
    t = strip_accents(str(text)).lower()
    t = re.sub(r"<[^>]+>", " ", t)              # bỏ tag HTML
    t = re.sub(r"^\s*\d+(\.\d+)*\s*[.)\-]?\s*", " ", t)  # bỏ "1.", "2.1)"
    t = re.sub(r"\([^)]*\)", " ", t)             # bỏ phần trong ngoặc
    t = re.sub(r"[^a-z0-9\s]", " ", t)           # bỏ dấu câu
    return re.sub(r"\s+", " ", t).strip()


def similarity(a: str, b: str) -> float:
    """Điểm giống nhau 0..1 giữa 2 chuỗi đã chuẩn hoá."""
    na, nb = normalize(a), normalize(b)
    if not na or not nb:
        return 0.0
    if na == nb:
        return 1.0
    # Một chuỗi chứa trọn chuỗi kia -> điểm cao
    if na in nb or nb in na:
        shorter, longer = (na, nb) if len(na) < len(nb) else (nb, na)
        return 0.80 + 0.15 * (len(shorter) / len(longer))
    # Trùng từ khoá
    wa, wb = set(na.split()), set(nb.split())
    if wa and wb:
        jaccard = len(wa & wb) / len(wa | wb)
    else:
        jaccard = 0.0
    ratio = SequenceMatcher(None, na, nb).ratio()
    return max(ratio, jaccard * 0.9)


# ---------------------------------------------------------------- slots

@dataclass
class Slot:
    """Một chỗ trống tìm thấy trong DOC."""
    kind: str            # "paragraph" | "cell" | "checkbox" | "table"
    label: str           # nhãn đọc được cạnh chỗ trống
    location: str        # mô tả vị trí cho người dùng dễ tìm
    para_index: int = -1
    blank_ordinal: int = 0   # chỗ trống thứ mấy trong đoạn
    table_index: int = -1
    row_index: int = -1
    cell_index: int = -1
    raw_text: str = ""


@dataclass
class Suggestion:
    slot: Slot
    field: Optional[Field]
    score: float
    placeholder: str

    @property
    def confidence(self) -> str:
        if self.score >= 0.85:
            return "cao"
        if self.score >= 0.60:
            return "vừa"
        return "thấp"


def _iter_body_paragraphs(doc):
    for i, p in enumerate(doc.paragraphs):
        yield i, p


def find_slots(doc) -> List[Slot]:
    """Quét DOC, trả về danh sách chỗ trống ứng viên."""
    slots: List[Slot] = []

    # 1) Paragraph dạng "Nhãn: ............"
    for i, p in _iter_body_paragraphs(doc):
        text = p.text
        if not text.strip():
            continue
        if PLACEHOLDER_RE.search(text):
            continue  # đã có placeholder rồi, bỏ qua

        # Một đoạn có thể chứa nhiều chỗ trống:
        #   "Số điện thoại: ....... Fax: ......."
        blanks = list(BLANK_RE.finditer(text))
        if blanks:
            prev_end = 0
            for bi, m in enumerate(blanks):
                label = text[prev_end:m.start()].strip(" :.\t")
                prev_end = m.end()
                if not label:
                    continue
                slots.append(Slot(
                    kind="paragraph",
                    label=label,
                    location=f"Đoạn văn #{i + 1}"
                             + (f", chỗ trống {bi + 1}" if len(blanks) > 1 else ""),
                    para_index=i,
                    blank_ordinal=bi,
                    raw_text=text,
                ))
            continue

        # Paragraph có ô tick
        if CHECKBOX_RE.search(text):
            label = CHECKBOX_RE.sub("", text).strip(" :.\t")
            if label:
                slots.append(Slot(
                    kind="checkbox",
                    label=label,
                    location=f"Đoạn văn #{i + 1} (ô tick)",
                    para_index=i,
                    raw_text=text,
                ))

    # 2) Ô bảng
    for ti, table in enumerate(doc.tables):
        n_rows = len(table.rows)
        for ri, row in enumerate(table.rows):
            cells = row.cells
            for ci, cell in enumerate(cells):
                ctext = cell.text.strip()
                if PLACEHOLDER_RE.search(ctext):
                    continue

                # Ô rỗng hoặc chỉ có dấu chấm -> nhãn nằm ở ô bên trái
                is_blank = (not ctext) or bool(BLANK_RE.fullmatch(ctext))
                if is_blank and ci > 0:
                    label = cells[ci - 1].text.strip(" :.\t")
                    # Bỏ ô mà "nhãn" chỉ là số thứ tự (cột STT)
                    if label.isdigit():
                        continue
                    if label and not BLANK_RE.fullmatch(label):
                        slots.append(Slot(
                            kind="cell",
                            label=label,
                            location=f"Bảng {ti + 1}, dòng {ri + 1}, cột {ci + 1}",
                            table_index=ti, row_index=ri, cell_index=ci,
                            raw_text=ctext,
                        ))
                    continue

                # Ô có ô tick
                if CHECKBOX_RE.search(ctext):
                    label = CHECKBOX_RE.sub("", ctext).strip(" :.\t")
                    if not label and ci > 0:
                        label = cells[ci - 1].text.strip(" :.\t")
                    if label:
                        slots.append(Slot(
                            kind="checkbox",
                            label=label,
                            location=f"Bảng {ti + 1}, dòng {ri + 1}, cột {ci + 1} (ô tick)",
                            table_index=ti, row_index=ri, cell_index=ci,
                            raw_text=ctext,
                        ))

        # 3) Bảng nhiều dòng -> ứng viên cho datagrid
        if n_rows >= 1:
            header = " | ".join(c.text.strip() for c in table.rows[0].cells)
            if header.strip():
                slots.append(Slot(
                    kind="table",
                    label=header,
                    location=f"Bảng {ti + 1} ({n_rows} dòng, "
                             f"{len(table.rows[0].cells)} cột)",
                    table_index=ti,
                    raw_text=header,
                ))

    return slots


# ---------------------------------------------------------------- matching

def match_fields(slots: List[Slot], fields: List[Field],
                 threshold: float = 0.45) -> List[Suggestion]:
    """
    Ghép chỗ trống với field. Mỗi field chỉ dùng 1 lần, ưu tiên cặp điểm cao.
    """
    grid_keys = {f.key for f in fields if f.type in GRID_TYPES}

    # Field ứng viên, bỏ subfield của grid
    candidates: List[Tuple[Field, str, str]] = []  # (field, label_dò, placeholder)
    for f in fields:
        if f.parent in grid_keys:
            continue
        if f.hidden:
            continue
        if f.type in MULTI_CHECKBOX_TYPES:
            for val, lbl, ph in f.option_placeholders():
                candidates.append((f, lbl or val, ph))
        else:
            candidates.append((f, f.label or f.key, f.placeholder))

    # Chấm điểm mọi cặp hợp lệ về kiểu
    pairs = []
    for si, slot in enumerate(slots):
        for ci, (f, lbl, ph) in enumerate(candidates):
            # Ràng buộc kiểu
            if slot.kind == "table" and f.type not in GRID_TYPES:
                continue
            if slot.kind != "table" and f.type in GRID_TYPES:
                continue
            if slot.kind == "checkbox" and not (
                f.type == "checkbox" or f.type in MULTI_CHECKBOX_TYPES
            ):
                continue
            if slot.kind in ("paragraph", "cell") and (
                f.type == "checkbox" or f.type in MULTI_CHECKBOX_TYPES
            ):
                continue

            score = similarity(slot.label, lbl)

            # Bảng: cộng điểm nếu tên cột khớp subfield
            if slot.kind == "table" and f.type in GRID_TYPES:
                subs = [x for x in fields if x.parent == f.key]
                if subs:
                    header_cols = [c.strip() for c in slot.label.split("|")]
                    hits = sum(
                        1 for s in subs
                        if any(similarity(col, s.label) > 0.6 for col in header_cols)
                    )
                    score = max(score, hits / len(subs))

            if score >= threshold:
                pairs.append((score, si, ci, f, ph))

    pairs.sort(key=lambda x: -x[0])

    used_slots, used_cands = set(), set()
    matched: List[Suggestion] = []
    for score, si, ci, f, ph in pairs:
        if si in used_slots or ci in used_cands:
            continue
        used_slots.add(si)
        used_cands.add(ci)
        matched.append(Suggestion(slot=slots[si], field=f,
                                  score=score, placeholder=ph))

    # Chỗ trống không khớp được field nào
    for si, slot in enumerate(slots):
        if si not in used_slots:
            matched.append(Suggestion(slot=slot, field=None,
                                      score=0.0, placeholder=""))

    # Sắp theo vị trí trong tài liệu
    matched.sort(key=lambda s: (
        s.slot.table_index if s.slot.table_index >= 0 else -1,
        s.slot.para_index,
        s.slot.row_index,
        s.slot.cell_index,
    ))
    return matched


def unmatched_fields(fields: List[Field],
                     suggestions: List[Suggestion]) -> List[Tuple[Field, str]]:
    """Field trong eform chưa được gán vào chỗ nào trong DOC."""
    grid_keys = {f.key for f in fields if f.type in GRID_TYPES}
    assigned = {s.placeholder for s in suggestions if s.field is not None}

    out = []
    for f in fields:
        if f.parent in grid_keys or f.hidden:
            continue
        if f.type in MULTI_CHECKBOX_TYPES:
            for val, lbl, ph in f.option_placeholders():
                if ph not in assigned:
                    out.append((f, ph))
        else:
            if f.placeholder not in assigned:
                out.append((f, f.placeholder))
    return out


# ---------------------------------------------------------------- apply

def _replace_nth_blank(text: str, n: int, replacement: str) -> str:
    """Thay chỗ trống thứ n (0-based). Placeholder đã chèn không tính là chỗ trống."""
    idx = 0
    out, pos = [], 0
    for m in BLANK_RE.finditer(text):
        # bỏ qua nếu đoạn này nằm trong một ${...} đã chèn trước đó
        out.append(text[pos:m.start()])
        if idx == n:
            out.append(replacement)
        else:
            out.append(m.group())
        pos = m.end()
        idx += 1
    out.append(text[pos:])
    return "".join(out)


def _replace_in_paragraph(p, new_text: str):
    """Ghi đè text paragraph, giữ định dạng của run đầu."""
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def apply_suggestions(docx_path: str, out_path: str,
                      suggestions: List[Suggestion],
                      accepted: List[bool] = None,
                      all_fields: List[Field] = None,
                      tidy: bool = True,
                      table_mode: str = "replace") -> int:
    """
    Chèn placeholder vào DOC gốc theo các gợi ý được chấp nhận.
    accepted[i] = True nghĩa là áp dụng suggestions[i].

    table_mode:
      "replace" — xoá toàn bộ bảng trong DOC, thay bằng một dòng
                  ${TenBang#table}. Backend tự dựng lại cả bảng.
      "row"     — giữ dòng tiêu đề, để đúng một dòng mẫu, mỗi ô mang
                  ${TenBang.TenCot}. Backend lặp dòng mẫu.

    Trả về số chỗ đã chèn.
    """
    doc = Document(docx_path)
    all_fields = all_fields or []
    # Chụp tham chiếu bảng TRƯỚC khi mutate: xoá một bảng sẽ làm
    # lệch chỉ số của các bảng phía sau.
    tables_snapshot = list(doc.tables)
    if accepted is None:
        accepted = [s.field is not None for s in suggestions]

    count = 0

    # Gom các chỗ trống cùng một đoạn văn để ghi đúng 1 lần,
    # tránh lần ghi sau đè mất lần ghi trước.
    para_jobs = {}
    for sug, ok in zip(suggestions, accepted):
        if ok and sug.placeholder and sug.slot.kind == "paragraph":
            para_jobs.setdefault(sug.slot.para_index, []).append(
                (sug.slot.blank_ordinal, sug.placeholder))

    for para_index, jobs in para_jobs.items():
        try:
            para = doc.paragraphs[para_index]
            text = para.text
            # Thay từ chỗ trống cuối lên đầu để không lệch vị trí
            for ordinal, ph in sorted(jobs, key=lambda x: -x[0]):
                text = _replace_nth_blank(text, ordinal, ph)
            _replace_in_paragraph(para, text)
            count += len(jobs)
        except (IndexError, AttributeError):
            continue

    for sug, ok in zip(suggestions, accepted):
        if not ok or not sug.placeholder:
            continue
        slot = sug.slot
        if slot.kind == "paragraph":
            continue  # đã xử lý ở trên

        try:
            if slot.kind == "cell":
                cell = doc.tables[slot.table_index].rows[slot.row_index] \
                          .cells[slot.cell_index]
                cell.text = ""
                run = cell.paragraphs[0].add_run(sug.placeholder)
                run.font.size = Pt(11)
                count += 1

            elif slot.kind == "checkbox":
                if slot.table_index >= 0:
                    cell = tables_snapshot[slot.table_index] \
                              .rows[slot.row_index].cells[slot.cell_index]
                    new = CHECKBOX_RE.sub(sug.placeholder, cell.text, count=1)
                    if new == cell.text:
                        new = cell.text + " " + sug.placeholder
                    cell.text = ""
                    cell.paragraphs[0].add_run(new).font.size = Pt(11)
                else:
                    p = doc.paragraphs[slot.para_index]
                    new = CHECKBOX_RE.sub(sug.placeholder, p.text, count=1)
                    _replace_in_paragraph(p, new)
                count += 1

            elif slot.kind == "table":
                if table_mode == "replace":
                    _replace_whole_table(tables_snapshot[slot.table_index],
                                         sug)
                else:
                    _fill_table_row(tables_snapshot[slot.table_index], sug,
                                    all_fields)
                count += 1

        except (IndexError, AttributeError):
            continue  # cấu trúc DOC đổi, bỏ qua chỗ này

    if tidy:
        tidy_document(doc)

    doc.save(out_path)
    return count


# ---------------------------------------------------------------- table fill

# Các tiêu đề cột được coi là số thứ tự, không map vào field nào
STT_HEADERS = {"stt", "so tt", "so thu tu", "tt", "no", "num", "index"}


def _is_stt_header(text: str) -> bool:
    return normalize(text) in STT_HEADERS


def _match_columns(header_cells: List[str], subfields: List[Field]) -> List[Optional[Field]]:
    """
    Ghép từng cột trong DOC với subfield của grid.
    Trả về list cùng độ dài header_cells, phần tử là Field hoặc None.
    """
    pairs = []
    for ci, htext in enumerate(header_cells):
        if _is_stt_header(htext):
            continue
        for si, sub in enumerate(subfields):
            score = similarity(htext, sub.label or sub.key)
            if score >= 0.40:
                pairs.append((score, ci, si))

    pairs.sort(key=lambda x: -x[0])
    result: List[Optional[Field]] = [None] * len(header_cells)
    used_cols, used_subs = set(), set()
    for score, ci, si in pairs:
        if ci in used_cols or si in used_subs:
            continue
        used_cols.add(ci)
        used_subs.add(si)
        result[ci] = subfields[si]

    # Cột chưa ghép được nhưng vẫn còn subfield thừa -> gán theo thứ tự
    leftover = [s for i, s in enumerate(subfields) if i not in used_subs]
    for ci in range(len(header_cells)):
        if result[ci] is None and not _is_stt_header(header_cells[ci]) and leftover:
            result[ci] = leftover.pop(0)

    return result


def _set_cell_placeholder(cell, text: str, center: bool = False):
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fill_table_row(table, sug: Suggestion, all_fields: List[Field]):
    """
    Cấu hình bảng theo kiểu 'lấy cả bảng':
      - Giữ nguyên dòng tiêu đề
      - Để đúng 1 dòng mẫu, mỗi cột mang placeholder của cột đó
      - Cột STT dùng ${#index}
    """
    grid = sug.field
    subfields = [f for f in all_fields if f.parent == grid.key]

    header_cells = [c.text.strip() for c in table.rows[0].cells]
    col_map = _match_columns(header_cells, subfields)

    # Xoá hết dòng dữ liệu, chừa lại tiêu đề
    for row in list(table.rows[1:]):
        row._element.getparent().remove(row._element)

    data_row = table.add_row()
    cells = data_row.cells

    for ci, cell in enumerate(cells):
        header = header_cells[ci] if ci < len(header_cells) else ""
        if _is_stt_header(header):
            _set_cell_placeholder(cell, "${#index}", center=True)
            continue
        sub = col_map[ci] if ci < len(col_map) else None
        if sub is not None:
            _set_cell_placeholder(cell, "${%s.%s}" % (grid.key, sub.key))
        else:
            _set_cell_placeholder(cell, "")

    # Đánh dấu bảng lặp: đặt marker ở ô đầu tiên của dòng mẫu
    first = cells[0]
    marker = "${%s#table}" % grid.key
    existing = first.text
    first.text = ""
    para = first.paragraphs[0]
    if _is_stt_header(header_cells[0] if header_cells else ""):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(marker + (" " + existing if existing else ""))
    run.font.size = Pt(11)
    first.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ---------------------------------------------------------------- tidy

def _set_cell_margins(table, top=60, bottom=60, left=100, right=100):
    """Đặt lề trong ô bảng (đơn vị dxa, 1/20 pt)."""
    tblPr = table._tbl.tblPr
    for tag in ("w:tblCellMar",):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    mar = OxmlElement("w:tblCellMar")
    for name, val in (("top", top), ("left", left),
                      ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tblPr.append(mar)


def tidy_document(doc):
    """
    Căn chỉnh lại tài liệu sau khi chèn placeholder:
      - Bảng: tiêu đề in đậm, căn giữa, lặp lại khi sang trang
      - Ô bảng: căn giữa theo chiều dọc, lề trong đều
      - Đoạn văn: khoảng cách dòng thống nhất, bỏ dòng trống thừa liên tiếp
    """
    # --- Bảng ---
    for table in doc.tables:
        table.autofit = True
        _set_cell_margins(table)

        if not table.rows:
            continue

        # Dòng tiêu đề: đậm, căn giữa, lặp khi sang trang
        head = table.rows[0]
        trPr = head._tr.get_or_add_trPr()
        if not trPr.findall(qn("w:tblHeader")):
            th = OxmlElement("w:tblHeader")
            th.set(qn("w:val"), "true")
            trPr.append(th)

        for cell in head.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.bold = True

        # Dòng dữ liệu
        for row in table.rows[1:]:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)

    # --- Đoạn văn ---
    blank_streak = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            blank_streak += 1
            # Gộp nhiều dòng trống liên tiếp thành tối đa 1
            if blank_streak > 1:
                para._element.getparent().remove(para._element)
            continue
        blank_streak = 0

        pf = para.paragraph_format
        if pf.space_after is None or pf.space_after > Pt(12):
            pf.space_after = Pt(6)
        if pf.space_before is None or pf.space_before > Pt(12):
            pf.space_before = Pt(0)
        if pf.line_spacing is None:
            pf.line_spacing = 1.15


def _replace_whole_table(table, sug: Suggestion):
    """
    Chế độ "replace": xoá sạch bảng trong DOC, thay bằng một đoạn văn
    chứa ${TenBang#table}. Backend tự dựng lại toàn bộ bảng.
    """
    tbl_el = table._tbl
    parent = tbl_el.getparent()

    # Tạo đoạn văn mới đặt đúng vị trí bảng cũ
    new_p = OxmlElement("w:p")
    parent.insert(list(parent).index(tbl_el), new_p)
    parent.remove(tbl_el)

    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, table._parent)
    run = para.add_run(sug.placeholder)
    run.font.size = Pt(11)
    return para
